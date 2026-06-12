"""Generate render/templates/reference.docx — the pandoc style template that
makes DOCX output look like the rendered PDF (Times, paper font sizes, A4 with
2.3cm margins, italic indented abstract, black bold headings, 8pt captions,
flush-left body text, full grid table borders).

Starts from pandoc's default reference.docx and restyles it with python-docx.
The result is committed; rebuilding is only needed when the style mapping here
changes. LaTeX size map (11pt article): \\LARGE=17, \\Large=14.5, \\large=12.
"""
from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

DEFAULT_OUT = Path(__file__).parent / "templates" / "reference.docx"
FONT = "Times New Roman"


def _strip_theme(rpr_owner, attrs=("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme")):
    """Remove theme font/color references so explicit values actually apply."""
    from docx.oxml.ns import qn

    rpr = rpr_owner.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is not None:
        for a in attrs:
            fonts.attrib.pop(qn(f"w:{a}"), None)
    color = rpr.find(qn("w:color"))
    if color is not None:
        color.attrib.pop(qn("w:themeColor"), None)


def _style(styles, style_id, *, name=None, size=None, bold=None, italic=None,
           black=False, align=None, left=None, right=None, first_line=None):
    from docx.shared import Pt, RGBColor

    st = styles[style_id]
    f = st.font
    if name is not None:
        f.name = name
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if black:
        f.color.rgb = RGBColor(0, 0, 0)
    _strip_theme(st.element)
    pf = getattr(st, "paragraph_format", None)
    if pf is not None:
        if align is not None:
            pf.alignment = align
        if left is not None:
            pf.left_indent = Pt(left)
        if right is not None:
            pf.right_indent = Pt(right)
        if first_line is not None:
            pf.first_line_indent = Pt(first_line)


def _add_table_borders(doc):
    """Full grid borders on the Table style — like the PDF's bordered tables."""
    from docx.oxml.ns import qn

    table_style = next(s for s in doc.styles if s.style_id == "Table")
    el = table_style.element
    tblpr = el.find(qn("w:tblPr"))
    if tblpr is None:
        tblpr = el.makeelement(qn("w:tblPr"), {})
        el.append(tblpr)
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is not None:
        tblpr.remove(borders)
    borders = tblpr.makeelement(qn("w:tblBorders"), {})
    tblpr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideV"):   # schema order
        edge = borders.makeelement(qn(f"w:{side}"), {})
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), "000000")
        borders.append(edge)


def build(out_path: Path | None = None) -> Path:
    """Write the styled reference.docx; returns its path."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Mm

    out_path = Path(out_path) if out_path else DEFAULT_OUT
    out_path.parent.mkdir(parents=True, exist_ok=True)

    default = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True, check=True,
    ).stdout
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(default)
        tmp = Path(f.name)
    try:
        doc = docx.Document(str(tmp))
    finally:
        tmp.unlink(missing_ok=True)

    styles = {s.style_id: s for s in doc.styles}
    C, J = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
    _style(styles, "Normal", name=FONT, size=11, black=True)
    _style(styles, "Title", name=FONT, size=17, bold=True, black=True, align=C)
    _style(styles, "Author", name=FONT, size=11, black=True, align=C)
    _style(styles, "Date", name=FONT, size=11, black=True, align=C)
    _style(styles, "AbstractTitle", name=FONT, size=14.5, bold=True,
           italic=False, black=True)
    _style(styles, "Abstract", name=FONT, size=11, italic=True, black=True,
           align=J, left=22, right=22)        # 2em at 11pt
    _style(styles, "Heading1", name=FONT, size=14.5, bold=True, italic=False,
           black=True)
    _style(styles, "Heading2", name=FONT, size=12, bold=True, italic=False,
           black=True)
    _style(styles, "Heading3", name=FONT, size=11, bold=True, italic=False,
           black=True)
    L = WD_ALIGN_PARAGRAPH.LEFT
    _style(styles, "BodyText", align=L, first_line=11)   # \parindent 1em, flush left
    _style(styles, "FirstParagraph", align=L, first_line=0)
    _style(styles, "ImageCaption", name=FONT, size=8, black=True, align=C)
    if "TableCaption" in styles:
        _style(styles, "TableCaption", name=FONT, size=8, black=True)
    _add_table_borders(doc)

    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)   # A4
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Cm(2.3))

    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUT
    print(f"Wrote {build(target)}")
