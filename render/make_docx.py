"""Markdown -> DOCX conversion for make-paper (post-render, on user request).

Runs the same preprocess pipeline as the PDF renderer but targets pandoc's
native docx writer: real heading styles, editable flowing text, PNG figures.
Tables are post-styled like the PDF (centered cells, shaded header row);
remaining LaTeX-only looks (dotted numbering) don't carry over — the DOCX is
the editable companion, the PDF stays the styled artifact.
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

import render_paper as rp


def docx_available() -> str | None:
    """None if pandoc is on PATH; else a warning with the install fix."""
    if shutil.which("pandoc") is None:
        return ("docx skipped (pandoc not on PATH — "
                "winget install JohnMacFarlane.Pandoc)")
    return None


_HEAD_STYLES = {"Title", "Author", "Date", "AbstractTitle", "Abstract"}
HEADER_FILL = "D9E2F3"   # the PDF's tableheadbg light blue


def style_tables(docx_path: Path) -> None:
    """Tables like the PDF: centered cells, light-blue shaded header row.

    Direct formatting (not a table style) because pandoc's Compact cell
    style would override conditional table-style formatting.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = docx.Document(str(docx_path))
    for table in doc.tables:
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if r == 0:
                    tcpr = cell._tc.get_or_add_tcPr()
                    shd = tcpr.find(qn("w:shd"))
                    if shd is None:
                        shd = tcpr.makeelement(qn("w:shd"), {})
                        tcpr.append(shd)
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), HEADER_FILL)
    doc.save(str(docx_path))


def apply_two_column_layout(docx_path: Path) -> None:
    """Two-column body with a full-width title head, like the PDF.

    Pandoc emits one single-column section; this closes the title block
    (Title/Author/Date/Abstract paragraphs) as its own one-column section and
    switches the rest to two columns continuing on the same page. A
    Literaturverzeichnis section at the end gets its own full-width section
    again, mirroring the PDF's \\onecolumn switch (rule v2.3-1).
    """
    from copy import deepcopy

    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(str(docx_path))
    body_sectpr = doc.sections[-1]._sectPr

    def set_cols(sectpr, num):
        cols = sectpr.find(qn("w:cols"))
        if cols is None:
            cols = sectpr.makeelement(qn("w:cols"), {})
            sectpr.append(cols)
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:space"), "454")        # 0.8cm, the PDF's \columnsep

    def set_continuous(sectpr):
        sect_type = sectpr.find(qn("w:type"))
        if sect_type is None:
            sect_type = sectpr.makeelement(qn("w:type"), {})
            sectpr.insert(0, sect_type)       # w:type precedes pgSz/pgMar
        sect_type.set(qn("w:val"), "continuous")

    last_head = None
    for p in doc.paragraphs:
        if p.style.style_id in _HEAD_STYLES:
            last_head = p
        else:
            break
    if last_head is not None:
        head_sectpr = deepcopy(body_sectpr)   # same page size/margins
        set_cols(head_sectpr, 1)
        last_head._p.get_or_add_pPr().append(head_sectpr)
        set_continuous(body_sectpr)

    refs_prev = prev = None
    for p in doc.paragraphs:
        if (p.style.style_id == "Heading1"
                and p.text.strip().lower() == "literaturverzeichnis"):
            refs_prev = prev
            break
        prev = p
    if refs_prev is not None:
        body_end = deepcopy(body_sectpr)      # closes the two-column body
        set_cols(body_end, 2)
        set_continuous(body_end)
        refs_prev._p.get_or_add_pPr().append(body_end)
        set_continuous(body_sectpr)
        set_cols(body_sectpr, 1)              # references: full width again
    else:
        set_cols(body_sectpr, 2)
    doc.save(str(docx_path))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="make-paper Markdown -> DOCX converter")
    parser.add_argument("--input", required=True, help="report Markdown from Subsystem A")
    parser.add_argument("--output", help="target DOCX path (default: input with .docx)")
    parser.add_argument("--project", required=True, help="project root (for images)")
    parser.add_argument("--title", help="override the title (else first H1)")
    parser.add_argument("--authors", help="author line, e.g. 'A 123, B 456'")
    parser.add_argument("--dateline", help="institution dateline, e.g. 'RWTH Aachen, Juni 2026'")
    parser.add_argument("--charts", choices=["auto", "blocks", "off"], default="blocks",
                        help="paperchart handling, same as the PDF renderer")
    parser.add_argument("--layout", choices=["one", "two"], default="one",
                        help="page layout, same as the PDF renderer")
    args = parser.parse_args(argv)

    missing = docx_available()
    if missing:
        print(missing, file=sys.stderr)
        return 2
    src = Path(args.input)
    if not src.is_file():
        print(f"input Markdown not found: {src}", file=sys.stderr)
        return 1

    out = rp.render(
        src, Path(args.output) if args.output else src.with_suffix(".docx"),
        args.project,
        title=args.title, authors=args.authors, dateline=args.dateline,
        charts=args.charts, to="docx",
    )
    if out.is_file():
        try:
            style_tables(out)
            if args.layout == "two":
                apply_two_column_layout(out)
        except ImportError as e:
            print(f"docx styling skipped ({e.name} not installed — "
                  "pip install python-docx)", file=sys.stderr)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
