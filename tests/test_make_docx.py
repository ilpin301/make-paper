import shutil
import sys
import zipfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "render"))

import make_docx as md


def test_docx_available_reports_missing_pandoc(monkeypatch):
    monkeypatch.setattr(md.shutil, "which", lambda name: None)
    assert "pandoc" in md.docx_available()
    monkeypatch.setattr(md.shutil, "which", lambda name: "/usr/bin/" + name)
    assert md.docx_available() is None


def test_main_errors_when_pandoc_missing(tmp_path, monkeypatch, capsys):
    monkeypatch.setattr(md, "docx_available", lambda: "docx skipped (pandoc not on PATH)")
    rc = md.main(["--input", str(tmp_path / "x.md"), "--project", str(tmp_path)])
    assert rc == 2
    assert "pandoc" in capsys.readouterr().err


def test_main_errors_when_input_missing(tmp_path, capsys):
    rc = md.main(["--input", str(tmp_path / "fehlt.md"), "--project", str(tmp_path)])
    assert rc == 1
    assert "fehlt.md" in capsys.readouterr().err


def test_main_renders_docx_next_to_input_by_default(tmp_path, monkeypatch, capsys):
    src = tmp_path / "Paper.md"
    src.write_text("# T\n\nText.\n", encoding="utf-8")
    seen = {}

    def fake_render(inp, outp, project, **kwargs):
        seen["args"] = (Path(inp), Path(outp), Path(project))
        seen["kwargs"] = kwargs
        return Path(outp)

    monkeypatch.setattr(md.rp, "render", fake_render)
    rc = md.main(["--input", str(src), "--project", str(tmp_path),
                  "--authors", "Anna Autor 1", "--dateline", "RWTH, Juni 2026"])
    assert rc == 0
    assert seen["args"] == (src, tmp_path / "Paper.docx", tmp_path)
    assert seen["kwargs"]["to"] == "docx"
    assert seen["kwargs"]["authors"] == "Anna Autor 1"
    assert seen["kwargs"]["charts"] == "blocks"
    assert "Paper.docx" in capsys.readouterr().out


def test_style_tables_centers_cells_and_shades_header(tmp_path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Kopf A"
    t.cell(0, 1).text = "Kopf B"
    t.cell(1, 0).text = "Wert"
    path = tmp_path / "t.docx"
    doc.save(str(path))

    md.style_tables(path)

    doc = docx.Document(str(path))
    t = doc.tables[0]
    for row in t.rows:
        for cell in row.cells:
            assert all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                       for p in cell.paragraphs)
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert xml.count(f'w:fill="{md.HEADER_FILL}"') == 2   # header cells only


needs_pandoc = pytest.mark.skipif(shutil.which("pandoc") is None,
                                  reason="pandoc not installed")


@needs_pandoc
def test_main_real_pandoc_produces_docx_with_styling(tmp_path):
    src = tmp_path / "Bericht.md"
    src.write_text(
        "# Stiltest\n\n### Abstract\nKurzfassung.\n\n"
        "### Einleitung\nKugelfallviskosimetrie von Glycerin.\n\n"
        "### Literaturverzeichnis\n1. Eine Quelle. Aachen.\n2. Zweite Quelle. Bonn.\n",
        encoding="utf-8",
    )
    rc = md.main(["--input", str(src), "--project", str(tmp_path),
                  "--authors", "Anna Autor 1", "--dateline", "RWTH, Juni 2026",
                  "--charts", "off"])
    assert rc == 0
    out = tmp_path / "Bericht.docx"
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "Kugelfallviskosimetrie" in xml
    assert "Anna Autor 1" in xml                  # author in title block
    assert "RWTH, Juni 2026" in xml               # dateline mapped to date
    assert "Literaturverzeichnis" in xml
