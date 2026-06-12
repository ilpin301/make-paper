import shutil
import sys
import zipfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "render"))

import build_reference_docx as brd
import make_docx as md

needs_pandoc = pytest.mark.skipif(shutil.which("pandoc") is None,
                                  reason="pandoc not installed")


@needs_pandoc
def test_build_reference_docx_applies_paper_styles(tmp_path):
    out = brd.build(tmp_path / "reference.docx")
    assert out.is_file()
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = docx.Document(str(out))
    styles = {s.style_id: s for s in doc.styles}

    assert styles["Normal"].font.name == "Times New Roman"
    assert styles["Normal"].font.size.pt == 11
    assert styles["Title"].font.size.pt == 17
    assert styles["Title"].font.bold is True
    assert styles["Title"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert styles["Heading1"].font.size.pt == 14.5
    assert styles["Heading1"].font.bold is True
    assert styles["Heading1"].font.name == "Times New Roman"
    assert styles["Heading2"].font.size.pt == 12
    assert styles["Abstract"].font.italic is True
    assert styles["Abstract"].paragraph_format.left_indent.pt == 22
    assert styles["ImageCaption"].font.size.pt == 8
    # v2.3 rule 1: body flush left (abstract stays justified)
    assert styles["BodyText"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert styles["FirstParagraph"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert styles["Abstract"].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    # v2.3 rule 4: table captions smaller than body text
    if "TableCaption" in styles:
        assert styles["TableCaption"].font.size.pt == 8

    section = doc.sections[0]
    assert round(section.left_margin.cm, 2) == 2.3
    assert round(section.page_width.cm, 1) == 21.0     # A4

    # v2.3 rule 4: full grid borders on the Table style (was insideV only)
    from docx.oxml.ns import qn
    table_style = next(s for s in doc.styles if s.style_id == "Table")
    borders = table_style.element.find(qn("w:tblPr")).find(qn("w:tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideV"):
        assert borders.find(qn(f"w:{side}")) is not None, side


def test_apply_two_column_layout_splits_head_and_body(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("Mein Titel", style="Title")
    doc.add_paragraph("Fliesstext des Berichts.")
    path = tmp_path / "d.docx"
    doc.save(str(path))

    md.apply_two_column_layout(path)

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    head, body = xml.split("</w:sectPr>")[0], xml.split("</w:sectPr>")[1]
    assert xml.count("<w:sectPr") == 2                 # head + body sections
    assert 'w:num="2"' not in head                     # head stays one column
    assert 'w:num="2"' in body                         # body is two columns
    assert 'w:val="continuous"' in body                # no page break between


def test_apply_two_column_layout_references_full_width(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("Mein Titel", style="Title")
    doc.add_paragraph("Fliesstext des Berichts.")
    doc.add_paragraph("Literaturverzeichnis", style="Heading 1")
    doc.add_paragraph("[1] Quelle Eins.")
    path = tmp_path / "d.docx"
    doc.save(str(path))

    md.apply_two_column_layout(path)

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    parts = xml.split("<w:sectPr")
    assert len(parts) == 4                  # head + body + references sections
    assert 'w:num="1"' in parts[1]          # head: one column
    assert 'w:num="2"' in parts[2]          # body: two columns
    assert 'w:num="1"' in parts[3]          # references: full width again
    assert 'w:val="continuous"' in parts[3]


@needs_pandoc
def test_docx_end_to_end_styled_two_column(tmp_path):
    src = tmp_path / "Bericht.md"
    src.write_text(
        "# Stiltest\n\n### Abstract\nKurzfassung des Berichts.\n\n"
        "### Einleitung\nKugelfallviskosimetrie von Glycerin.\n\n"
        "### Literaturverzeichnis\n1. Eine Quelle. Aachen.\n",
        encoding="utf-8",
    )
    rc = md.main(["--input", str(src), "--project", str(tmp_path),
                  "--authors", "Anna Autor 1", "--dateline", "RWTH, Juni 2026",
                  "--charts", "off", "--layout", "two"])
    assert rc == 0
    with zipfile.ZipFile(tmp_path / "Bericht.docx") as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
        styles_xml = z.read("word/styles.xml").decode("utf-8")
    assert "Times New Roman" in styles_xml             # reference.docx applied
    assert 'w:num="2"' in doc_xml                      # two-column body
    # full-width head + two-column body + full-width references (v2.3)
    assert doc_xml.count("<w:sectPr") == 3
